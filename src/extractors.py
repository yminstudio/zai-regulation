"""첨부파일 텍스트 추출기.

지원 형식:
- PDF (텍스트 우선, 부족하면 페이지 OCR)
- HWP/HWPX (가능 시 추출, 실패 시 빈 문자열)
- Excel (xls/xlsx)
- Word (docx/doc)
- 이미지 (OCR)
"""
from __future__ import annotations

import io
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

import pytesseract
import xlrd
from PIL import Image
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
import pypdfium2 as pdfium

OCR_LANG = "kor+eng"
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}


def _clean_text(value: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", (value or "").strip())


def _read_text_with_fallback(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return raw.decode(enc)
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace")


def _convert_with_office(input_path: Path, out_dir: Path, target: str) -> Path | None:
    """LibreOffice/soffice/lowriter 중 가능한 실행기로 변환."""
    for cmd in ("libreoffice", "soffice", "lowriter"):
        try:
            proc = subprocess.run(
                [
                    cmd,
                    "--headless",
                    "--convert-to",
                    target,
                    "--outdir",
                    str(out_dir),
                    str(input_path.resolve()),
                ],
                capture_output=True,
                timeout=60,
            )
            if proc.returncode != 0:
                continue
            produced = sorted(out_dir.glob(f"{input_path.stem}.*"))
            if produced:
                return produced[0]
        except Exception:
            continue
    return None


def _ocr_image_file(path: Path) -> str:
    img = Image.open(path)
    return _clean_text(pytesseract.image_to_string(img, lang=OCR_LANG))


def _extract_pdf_text(path: Path) -> tuple[str, str]:
    """PDF는 페이지별로 텍스트 품질을 보고 OCR을 혼합 적용."""
    reader = PdfReader(str(path))
    doc = pdfium.PdfDocument(str(path))
    parts: list[str] = []
    used_ocr = False

    for idx, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        # 페이지 텍스트가 충분하면 우선 사용
        if len(page_text) >= 80:
            parts.append(page_text)
            continue

        # 부족한 페이지만 OCR
        used_ocr = True
        pil = doc[idx].render(scale=2).to_pil()
        ocr_text = pytesseract.image_to_string(pil, lang=OCR_LANG).strip()
        if ocr_text:
            parts.append(ocr_text)
        else:
            parts.append(page_text)

    text = _clean_text("\n".join(parts))
    if used_ocr:
        return text, "pdf_mixed_ocr"
    return text, "pdf_text"


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return _clean_text("\n".join(lines))


def _extract_doc_text(path: Path) -> tuple[str, str]:
    """.doc 추출: antiword -> catdoc -> office(txt) -> office(docx) -> 변환필요."""
    # antiword
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        return _clean_text(result.stdout), "doc_antiword"
    except Exception:
        pass
    # catdoc
    try:
        result = subprocess.run(
            ["catdoc", str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        return _clean_text(result.stdout), "doc_catdoc"
    except Exception:
        pass
    # 방법 A-1: office로 .txt 변환 후 읽기
    try:
        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            txt_path = _convert_with_office(path, out_dir, "txt")
            if txt_path and txt_path.suffix.lower() == ".txt":
                text = _read_text_with_fallback(txt_path)
                return _clean_text(text), "doc_office_txt"
    except Exception:
        pass
    # 방법 A-2: office로 .docx 변환 후 python-docx로 읽기
    try:
        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            docx_path = _convert_with_office(path, out_dir, "docx")
            if docx_path and docx_path.suffix.lower() == ".docx":
                return _extract_docx_text(docx_path), "doc_office_docx"
    except Exception:
        pass
    # 실무 정책: .doc는 변환 성공 시에만 처리, 실패하면 변환 필요로 남김
    return "", "doc_needs_docx_conversion"


def _extract_xlsx_text(path: Path) -> str:
    wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"[Sheet] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if vals:
                out.append(" | ".join(vals))
    return _clean_text("\n".join(out))


def _extract_xls_text(path: Path) -> str:
    wb = xlrd.open_workbook(str(path))
    out: list[str] = []
    for sheet in wb.sheets():
        out.append(f"[Sheet] {sheet.name}")
        for r in range(sheet.nrows):
            vals = []
            for c in range(sheet.ncols):
                v = sheet.cell_value(r, c)
                if str(v).strip():
                    vals.append(str(v).strip())
            if vals:
                out.append(" | ".join(vals))
    return _clean_text("\n".join(out))


def _extract_hwpx_text(path: Path) -> str:
    # HWPX는 zip 기반 XML. 텍스트 노드만 단순 추출.
    try:
        with zipfile.ZipFile(path, "r") as zf:
            lines: list[str] = []
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                data = zf.read(name).decode("utf-8", errors="ignore")
                lines.extend(re.findall(r">([^<>]+)<", data))
            return _clean_text("\n".join(lines))
    except Exception:
        return ""


def _extract_hwp_text(path: Path) -> tuple[str, str]:
    """HWP 추출: hwp5txt -> LibreOffice headless -> strings. 반환 (텍스트, 방법)."""
    try:
        result = subprocess.run(
            ["hwp5txt", str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        return _clean_text(result.stdout), "hwp_hwp5txt"
    except Exception:
        pass
    try:
        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            txt_path = _convert_with_office(path, out_dir, "txt")
            if txt_path and txt_path.suffix.lower() == ".txt":
                text = _read_text_with_fallback(txt_path)
                return _clean_text(text), "hwp_office_txt"
    except Exception:
        pass
    try:
        result = subprocess.run(
            ["strings", str(path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        return _clean_text(result.stdout), "hwp_strings"
    except Exception:
        return "", "hwp_failed"


def extract_text(path: str | Path) -> tuple[str, str]:
    """파일 텍스트 추출.

    Returns:
        (text, method)
    """
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_text(p)
    if ext in IMAGE_EXTS:
        return _ocr_image_file(p), "image_ocr"
    if ext == ".docx":
        return _extract_docx_text(p), "docx"
    if ext == ".doc":
        return _extract_doc_text(p)
    if ext == ".xlsx":
        return _extract_xlsx_text(p), "xlsx"
    if ext == ".xls":
        return _extract_xls_text(p), "xls"
    if ext == ".hwpx":
        return _extract_hwpx_text(p), "hwpx"
    if ext == ".hwp":
        return _extract_hwp_text(p)

    # 텍스트 계열 fallback
    try:
        data = Path(path).read_text(encoding="utf-8", errors="ignore")
        return _clean_text(data), "plain_text"
    except Exception:
        return "", "unsupported"

